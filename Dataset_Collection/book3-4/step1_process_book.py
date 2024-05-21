from docx import Document
from collections import defaultdict
import re
import csv


def match_book3(input_path):
    document = Document(input_path)
    cn_result = defaultdict(str)
    en_result = defaultdict(str)

    pattern_cn_en = r'^(图([\d-]*).*)(Fig.*)'
    pattern_cn = r'^图([\d-]*).*'
    pattern_en = r'Fig[\. ]*([\d-]*).*'
    # pattern_en = r'FIGURE[\. ]*([\d-]*).*' #book4
    pattern_multi2one = r'[A-Z][\.．]'

    multipic_count = 0
    for para in document.paragraphs:
        match = re.search(pattern_cn_en, para.text)

        # Match according to the pattern "pattern_cn_en", assuming that this paragraph contains both a cn_caption and an en_caption.
        if match:
            # print("1_match",match)
            id = match.group(2)
            cn_result[id] = match.group(1)
            en_result[id] = match.group(3)
            is_multipic_match = re.search(pattern_multi2one, match.group(1))
            if (is_multipic_match):
                print(match.group(1))

            continue
        
        # If the paragraph contains only a cn_caption or only an en_caption, then match once according to pattern_cn or pattern_en, respectively.
        match = re.search(pattern_cn, para.text)
        if match:
            # print("2_match",match)
            id = match.group(1)
            cn_result[id] = match.string
            is_multipic_match = re.search(pattern_multi2one, match.string)
            if (is_multipic_match):
                print(match.string)
                multipic_count += 1
            continue
        
        match = re.search(pattern_en, para.text)
        if match:
            # print("3_match",match)
            id = match.group(1)
            en_result[id] = match.string

    print("en_result",en_result)
    print("cn_result",cn_result)
    return cn_result, en_result, multipic_count

def match_book4(input_path):
    document = Document(input_path)
    cn_result = defaultdict(str)
    en_result = defaultdict(str)

    # pattern_cn_en = r'^(图([\d-]*).*)(Fig.*)'
    # pattern_cn = r'^图([\d-]*).*'
    # pattern_en = r'Fig[\. ]*([\d-]*).*' #book3
    pattern_en = r'FIGURE\s\d+-\d+\.' #book4
    pattern_multi2one = r'A[\.．]'

    multipic_count = 0
    for para in document.paragraphs:

        
        match = re.search(pattern_en, para.text)
        print("match",match)
        if match:
            print("match.group(0)",match.group(0))
            id = match.group(0)
            en_result[id] = match.string

    print("en_result",en_result)
    # print("cn_result",cn_result)
    return en_result, multipic_count


def write_into_csv_book4(output_path, match_info):
    en_result, multipic_count = match_info
    pattern_multi2one = r'[A-Z][\.．]'
    pattern = r'[A-Z]'

    with open(output_path, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Image_ID", "cn_caption", "en_caption", "is_multipic"])

        for key in en_result.keys():
            # print("cn_result[key]",cn_result[key])
            m_match = re.search(pattern_multi2one, en_result[key])
            is_multipic = ''
            # print("m_match",m_match)
            if (m_match):
                is_multipic = 'Y'
            u_match = re.search(pattern, key)
            if u_match:
                name = key.split(" ")[-1].split('.')[0]
            else:
                name = key
            writer.writerow(['figure'+name, '',
                            en_result[key], is_multipic])

def write_into_csv_book3(output_path, match_info):
    cn_result,en_result, multipic_count = match_info
    pattern_multi2one = r'[A-Z][\.．]'


    with open(output_path, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Image_ID", "cn_caption", "en_caption", "is_multipic"])

        for key in cn_result.keys():
            # print("cn_result[key]",cn_result[key])
            m_match = re.search(pattern_multi2one, cn_result[key])
            is_multipic = ''
            # print("m_match",m_match)
            if (m_match):
                is_multipic = 'Y'
            writer.writerow(['figure'+key, cn_result[key],
                            en_result[key], is_multipic])
         
if __name__ == "__main__":
    book_path = './book4/'  #replace
    total_part_number = 1   #replace (the number of sub-PDFs)
    for i in range(total_part_number):
        input_docx_path = book_path + 'texts/part' + str(i+1) + '.docx'
        output_csv_path = book_path + 'caption/part' + str(i+1)  + '.csv'
        match_result = match_book4(input_docx_path)         #replace
        write_into_csv_book4(output_csv_path, match_result) #replace
