import os
import csv
import re
from typing import List, Dict
from docx import Document
import docx
from collections import defaultdict
import binascii
from PIL import Image
import pandas as pd


def delete_file(file_path):
    if os.path.isfile(file_path):
        os.remove(file_path)
    else:
        print(f"Error: {file_path} not a valid filename")

# Extract text from paragraphs in docx
# extract_docx_book1 is used for extracting book1
def extract_docx_book1(docx_file_path):
    document = Document(docx_file_path)
    paragraphs_list = document.paragraphs
    print("paragraphs_list.length", len(paragraphs_list))

    # Patterns for matching: caption with both Chinese and English, caption with only Chinese, caption with only English
    # The following are for book3
    pattern_cn_en = r'^(图([\d-]*).*)(Fig.*)'
    pattern_cn = r'^图([\d-]*).*'
    pattern_en = r'Fig[\. ]*([\d-]*).*'
    pattern_is_multi2one = r'A[\.．]'
 
    # Chinese captions and English captions are stored in different Dicts
    cn_caption: Dict[str, str] = defaultdict(str)
    en_caption: Dict[str, str] = defaultdict(str)

    id = 0
 
    for i in range(len(paragraphs_list)):
        str_judge = ''
        # cn_en
        match = re.search(pattern_cn_en, paragraphs_list[i].text)
        if match:
            print("paragraphs_list[i].text", paragraphs_list[i].text)
            id = match.group(2)
            cn_caption[id] = match.group(1)
            en_caption[id] = match.group(3)
            m_match = re.search(pattern_is_multi2one, match.group(1))
            if(m_match):
                print(match.group(1))
                count += 1
                continue

        # cn
        match = re.search(pattern_cn, paragraphs_list[i].text)
        if match:
            id = match.string.split('\t')[0].replace('图', '')
            cn_caption[str(id)] = match.string
            m_match = re.search(pattern_is_multi2one, match.string)
            if(m_match):
                print(match.string)
            continue

        # en
        match = re.search(pattern_en, paragraphs_list[i].text)
        if match:
            id = match.group(1)
            en_caption[id] = match.string

    print("cn_caption", cn_caption)
    print("len(cn_caption)", len(cn_caption))
    return cn_caption, en_caption

# extract_docx_book2 is used for extracting book2
def extract_docx_book2(docx_file_path):
    temp_caption = ''
    document = Document(docx_file_path)
    paragraphs_list = document.paragraphs
    fake_id = 0
    id = None
    flag = False

    caption: Dict[str, str] = defaultdict(str)
    pattern_id = r'(图 *|Fig\.|FIGURE )'
    for i in range(len(paragraphs_list)):
        str_para: str = clear_space(paragraphs_list[i].text, 1)
        match = re.search(pattern_id, str_para)
        if match:
            id = match.string.split('\t')[0].replace('图', '')
            caption[id] = str_para
            fake_id += 1
        elif id != None:
            caption[id] += str_para
        else:
            pass
    print("caption", caption)
    return caption

# extract_docx_book4 is used for extracting book4
def extract_docx_book4(docx_file_path):
    document = Document(docx_file_path)
    paragraphs_list = document.paragraphs
    print("paragraphs_list.length", len(paragraphs_list))

    # Patterns for matching: caption with both Chinese and English, caption with only Chinese, caption with only English
    # The following are for book3
    pattern_cn_en = r'^(图([\d-]*).*)(Fig.*)'
    pattern_cn = r'^图([\d-]*).*'
    pattern_en = r'FIGURE[\. ]*([\d-]*).*'
    pattern_is_multi2one = r'A[\.．]'
 
    # Chinese captions and English captions are stored in different Dicts
    cn_caption: Dict[str, str] = defaultdict(str)
    en_caption: Dict[str, str] = defaultdict(str)

    id = 0
 
    for i in range(len(paragraphs_list)):
        str_judge = ''
        # cn_en
        match = re.search(pattern_cn_en, paragraphs_list[i].text)
        if match:
            print("paragraphs_list[i].text", paragraphs_list[i].text)
            id = match.group(2)
            cn_caption[id] = match.group(1)
            en_caption[id] = match.group(3)
            m_match = re.search(pattern_is_multi2one, match.group(1))
            if(m_match):
                print(match.group(1))
                count += 1
                continue

        # cn
        match = re.search(pattern_cn, paragraphs_list[i].text)
        if match:
            print("paragraphs_list[i].text", paragraphs_list[i].text)
            print("match.string", match.string)
            id = match.string.split('\t')[0].replace('图', '')
            cn_caption[str(id)] = match.string
            m_match = re.search(pattern_is_multi2one, match.string)
            if(m_match):
                print(match.string)
            continue

        # en
        match = re.search(pattern_en, paragraphs_list[i].text)
        if match:
            id = match.group(1)
            en_caption[id] = match.string

    print("cn_caption", cn_caption)
    print("len(en_caption)", len(en_caption))
    return cn_caption, en_caption

# Clear strange spaces in the text
def clear_space(content: str, choice: int):
    hex_content = binascii.hexlify(content.encode())
    print(hex_content)
    if(choice == 1):
        new_hex_content = hex_content.replace(b'0a0d', b'')

    new_content = binascii.unhexlify(new_hex_content)
    with open('filename.txt', 'wb') as file:
        file.write(new_content)
    return new_content.decode(errors='ignore')

def write_into_csv_book1(csv_path, pair_path, raw_caption):
    
    pattern_is_multi2one = r'[A-Z]{2,}'
 
    with open(csv_path, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Image_ID", "cn_caption", "en_caption", "is_multipic"])
        
        pair_files = os.listdir(pair_path)
        pair_files = [file for file in pair_files if file.endswith(
            ('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'))]
        pair_files = sorted(pair_files)
        # For book1
        cn_caption, en_caption = raw_caption

        for i, key in enumerate(cn_caption.keys()):
            print("cn_caption[k]", cn_caption[key])
            m_match = re.search(pattern_is_multi2one, pair_files[i])
            is_multipic = ''
            if(m_match):
                is_multipic = 'Y'
            writer.writerow([pair_files[i], cn_caption[key], en_caption[key], is_multipic])

    return file

def write_into_csv_book2(csv_path, pair_path, raw_caption):
    
    pattern_is_multi2one = r'[A-Z]{2,}'
    pattern_caption = r'^图'
    
    with open(csv_path, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Image_ID", "cn_caption", "en_caption", "is_multipic"])
        
        pair_files = os.listdir(pair_path)
        pair_files = [file for file in pair_files if file.endswith(
            ('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'))]
        pair_files = sorted(pair_files)

        caption = raw_caption
        print("caption", len(caption))
        
        i = 0
        for key in caption.keys():
            m_match = re.search(pattern_is_multi2one, pair_files[i])
            is_multipic = ''
            if(m_match):
                is_multipic = 'Y'
            c_match = re.search(pattern_caption, caption[key])
            if(c_match):
                writer.writerow([pair_files[i], caption[key], '', is_multipic])
                i = i + 1
            else:
                continue
        
    return file


def remove_blank_image(cur_path, pictures_list: List[str]):
    print(f"The initial list of pictures: {pictures_list}")
    for picture in pictures_list:
        if(picture.find('0001') != -1):
            pictures_list.remove(picture)
            delete_file(cur_path + picture)
            print(f"{picture} is removed.")

def delete_file(file_path):
    if os.path.isfile(file_path):
        os.remove(file_path)
    else:
        print(f"Error: {file_path} not a valid filename")
        
    print("done")

def rename_image(folder_path,new_folder_path,new_name: str):
    global file_index
    if file_index > len(image_files_list):
        print("太多了！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！")
        exit()
    old_name = image_files_list[file_index]
    print(f"{file_index},file_index;\n{new_name},new name")
    print("old_name",old_name)
    pic_org = Image.open(os.path.join(folder_path, old_name))  
    picture_path = os.path.join(new_folder_path, new_name)
    print("picture_path",picture_path)
    pic_org.save(picture_path)
    # file_index += 1
 
def rename_all_images(folder_path,new_folder_path,df):
    global file_index
    reference_list: List = list( zip(df['Image_ID'], df['cn_caption'], df['en_caption'], df['is_multipic']))
    result = []
    for item in reference_list:
        print("\nitem",item)
        if (item[3] == 'Y'):
            pattern_is_multi2one = r'[A-Z]{2,}'
            match = re.search(pattern_is_multi2one, str(item[0]))
            if (match):
                print("match",match)
                # while (match):
                upperChar = [char for char in str(item[0]) if char.isupper()]
                print("upperChar",upperChar)
                no_uppercase_image_name = [char for char in str(item[0]) if not char.isupper()]
                no_uppercase_image_name = ''.join(no_uppercase_image_name)
                print("no_uppercase_image_name",no_uppercase_image_name)
                for c in upperChar:
                    print("new_image_name",no_uppercase_image_name.split('.')[0]+c+'.jpg')
                    new_image_name = no_uppercase_image_name.split('.')[0]+c+'.jpg'
                    rename_image(folder_path,new_folder_path,new_image_name)
                    temp = (new_image_name, item[1], item[2], item[3])
                    result.append(temp)
        else:
            rename_image(folder_path,new_folder_path,item[0])
            result.append(item)
        file_index += 1
    print("result",result)
    return result
    
if(__name__=="__main__"):
    file_index = 0
     
    path = '../book1_pair_zcr/CFP/'              # replace 
    pair_path = '../book1_pair_zcr/CFP_pair/'    # replace
    
    input_docx_file= path + 'output_CFP.docx'    # replace     
    image_path = path + 'images_raw/'            
    image_new_path = path + 'images/'            
    csv_path = path + 'image_caption_book1.csv'  # replace
    output_path = path + 'image_caption_book1_result.csv'  # replace
    if(not os.path.exists(image_new_path)):
        os.mkdir(image_new_path)
        
    all_files_list = os.listdir(image_path)
    image_files_list = [file for file in all_files_list if file.endswith(
        ('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'))]
    image_files_list = sorted(image_files_list)
    print("len(image_files_list)",len(image_files_list))

    remove_blank_image(image_path,image_files_list)

    write_into_csv_book1(csv_path,pair_path,extract_docx_book1(input_docx_file))  # replace (replace function extract_docx_book1 and function write_into_csv_book1)
    
    df = pd.read_csv(csv_path, encoding='utf-8')        
    result = rename_all_images(image_path,image_new_path,df)    
    with open(output_path, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Image_ID", "cn_caption", "en_caption", "is_multipic"])
        for item in result:
            writer.writerow([item[0], item[1], item[2], item[3]])
